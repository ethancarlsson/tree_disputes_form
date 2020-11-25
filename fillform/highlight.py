import docx
from docx.table import _Cell
from docx.styles.style import _TableStyle
import os
from pathlib import Path
from io import BytesIO


def Write1(csrf, q1, q2, q3, q4):
    dir_path = str(os.path.dirname(os.path.realpath(__file__)))
    doc = docx.Document(dir_path + r'/forms/Tree Dispute Application Form C Version 3.docx')
    
    def ReadDocument(table, row, column):
        return doc.tables[table].cell(row, column)
    def WriteDocument(table, row, column):
        ReadDocument(table, row, column).text = "*Fill this out"
    ## Always highlighted ##
    WriteDocument(0,5,1)
    WriteDocument(0,6,1)
    WriteDocument(0,7,1)
    WriteDocument(0,8,1)

    WriteDocument(0,14,1)
    WriteDocument(0,16,1)
    WriteDocument(0,17,1)
    WriteDocument(0,18,1)
    WriteDocument(0,19,1)

    # full name of all owners of the property
    WriteDocument(0,20,1)
    ReadDocument(0,20,1).add_paragraph('1.   ………………………………………………………', style='heading 9')
    ReadDocument(0,20,1).add_paragraph('2.   ………………………………………………………', style='heading 9')
    ReadDocument(0,20,1).add_paragraph('3.   ………………………………………………………', style='heading 9')

    ##
    WriteDocument(0,22,1)

    if q1 == "Yes I will have a legal representative.":
        WriteDocument(0,10,1)
        ReadDocument(0,10,1).add_paragraph("[solicitor on record] [firm]", style='heading 9')
    elif q1 == "Yes I will have an agent representing me.":
        WriteDocument(0,11,1)
    elif q1 == "Yes I will have an authorised officer representing me.":
        WriteDocument(0,12,1)
    else:
        pass

    if q2 == "No, the owner lives somewhere else.":
        WriteDocument(0,21,1)
        WriteDocument(0,23,1)
        WriteDocument(0,24,1)
        ReadDocument(4,3,0).add_paragraph("*This applies to you")
    else:
        ReadDocument(0,21,1).text = "same as above"
        ReadDocument(4,5,0).add_paragraph("*This applies to you")
    
    if q3 == "It is a tree that could injure someone." or q3 == "It is a tree that could damage, has damaged or is damaging my property.":
        doc.paragraphs[10].text = "*Fill this out"
    else: 
        doc.paragraphs[25].text = "*Fill this out"

    if q4 == "Yes":
        ReadDocument(4,7,0).add_paragraph("*This applies to you")
    else:
        pass
    unique = csrf[5:15]
    root_folder = str(Path(dir_path).parent)
    doc.save(root_folder + fr'/formfolder/Tree Dispute Application Form C Version 3{unique}.docx')

def Write2(csrf, q1):
    dir_path = str(os.path.dirname(os.path.realpath(__file__)))
    if q1 == "It is a tree that could injure someone.":
        doc = docx.Document(dir_path + r'/forms/application_treedispute_damagetoproperty.docx')
    elif q1 == "It is a tree that could damage, has damaged or is damaging my property.":
        doc = docx.Document(dir_path + r'/forms/application_treedispute_damagetoproperty.docx')
    else:
        doc = docx.Document(dir_path + r'/forms/application_treedispute_claimdetails_highhedges.docx')
        

    if q1 == "It is a tree that could injure someone.":
        doc.paragraphs[8].runs[9].text = '*This applies to you'
        doc.paragraphs[8].runs[9].underline = True  
    elif q1 == "It is a tree that could damage, has damaged or is damaging my property.":
        doc.paragraphs[6].runs[8].text = '*This applies to you'
        doc.paragraphs[6].runs[8].underline = True 
    elif q1 == "It is a hedge that is obstructing sunlight to a window on my property.":
        doc.paragraphs[6].runs[9].text = '*This applies to you'
        doc.paragraphs[6].runs[9].underline = True  
    elif q1 == "It is a hedge that is obstructing a view from my property.":
        doc.paragraphs[8].runs[10].text = ' *This applies to you'
        doc.paragraphs[8].runs[10].underline = True

    unique = csrf[5:15]
    root_folder = str(Path(dir_path).parent)
    if q1 == "It is a tree that could injure someone.":
        doc.save(root_folder + fr'/formfolder/application_treedispute_damage_to_property{unique}.docx')
    elif q1 == "It is a tree that could damage, has damaged or is damaging my property.":
        doc.save(root_folder + fr'/formfolder/application_treedispute_damage_to_property{unique}.docx')
    else:
        doc.save(root_folder + fr'/formfolder/application_treedispute_claimdetails_highhedges{unique}.docx')